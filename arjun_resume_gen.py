"""
JD-Specific Resume Generator — Arjun Kumar
============================================
Reads arjun_jobs_top.csv (from JobSpy), sends each JD + base resume
to Claude API, and writes a tailored .docx per company.

Setup:
    pip install requests pandas python-docx
    export ANTHROPIC_API_KEY="sk-ant-..."

Run:
    python arjun_resume_gen.py
    python arjun_resume_gen.py --csv my_jobs.csv   # custom CSV
    python arjun_resume_gen.py --limit 5            # process only 5 jobs
    python arjun_resume_gen.py --min-score 60       # skip jobs below 60%

Output:
    output_resumes/
        Razorpay_Senior_Backend_Engineer.docx
        Zomato_Full_Stack_Engineer.docx
        ...
        summary.csv   <- score, company, file, apply URL
"""

import os, re, csv, time, json, argparse, textwrap
import requests
import pandas as pd
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Base resume data ──────────────────────────────────────────────────────────

BASE_RESUME = {
    "name": "Arjun Kumar",
    "title": "Senior Software Engineer | Full Stack & Backend",
    "contact": {
        "location": "Delhi, India",
        "phone": "+91 8949342705",
        "email": "kumar.arjun000511@gmail.com",
        "linkedin": "linkedin.com/in/arjun-kumar-b4150b10b",
    },
    "summary": (
        "Senior Software Engineer with 7+ years of experience designing, developing, "
        "and deploying scalable enterprise and SaaS applications using Node.js, NestJS, "
        "Angular, and AWS. Strong expertise in backend architecture, REST API development, "
        "authentication systems (JWT/RBAC), CI/CD automation, and cloud deployments. "
        "Experienced in building scalable platforms across EdTech, SportsTech, and B2B "
        "domains with focus on performance optimization, system reliability, and secure "
        "application architecture."
    ),
    "skills": {
        "Backend":         "Node.js, NestJS, Express.js, REST APIs, JWT, RBAC",
        "Frontend":        "Angular, TypeScript, HTML, CSS",
        "Database":        "MySQL, Redis, MongoDB",
        "Cloud & DevOps":  "AWS, CI/CD Pipelines, Git, Deployment Automation",
        "Architecture":    "System Design, Microservices, Scalable Systems",
        "Integrations":    "Stripe, Google Maps SDK, Third-party APIs",
    },
    "experience": [
        {
            "role": "Software Engineer",
            "company": "Quokkalabs LLP",
            "period": "Feb 2022 – Present",
            "bullets": [
                "Engineered scalable full-stack applications using Angular, Node.js, and NestJS for enterprise-grade SaaS platforms.",
                "Designed and developed secure REST APIs with JWT authentication and RBAC authorization mechanisms.",
                "Integrated Stripe payment gateway, Google Maps SDK, and multiple third-party APIs to enhance platform capabilities.",
                "Managed AWS-based production deployments and CI/CD pipelines, improving deployment efficiency and release stability.",
                "Optimized backend APIs and database queries, improving application responsiveness and overall system performance.",
                "Reduced deployment time by approximately 20% through automation and CI/CD improvements.",
            ],
        },
        {
            "role": "Software Engineer",
            "company": "InsightGeeks Solutions Pvt Ltd",
            "period": "Aug 2019 – Nov 2021",
            "bullets": [
                "Developed Angular-based web applications with reusable UI components and scalable frontend architecture.",
                "Integrated frontend systems with backend REST APIs ensuring seamless communication and data flow.",
                "Improved UI/UX and optimized frontend performance, enhancing user engagement and application responsiveness.",
                "Implemented responsive and cross-browser compatible designs for enterprise applications.",
            ],
        },
        {
            "role": "Frontend Developer",
            "company": "Elroute Pvt Ltd",
            "period": "Jun 2018 – Jul 2019",
            "bullets": [
                "Developed responsive Angular UI modules for enterprise applications.",
                "Implemented reusable frontend components to improve maintainability and consistency.",
                "Integrated frontend modules with backend APIs ensuring smooth application workflows.",
            ],
        },
    ],
    "projects": [
        {
            "name": "LevelExam (EdTech Platform)",
            "bullets": [
                "Developed scalable REST APIs using Node.js and Express for educational workflows.",
                "Integrated Moodle Web Services for content delivery and learning management.",
                "Implemented JWT authentication and RBAC-based authorization for secure access management.",
                "Managed AWS deployments and CI/CD pipelines for production releases.",
            ],
        },
        {
            "name": "Edique Control Panel",
            "bullets": [
                "Built backend APIs using NestJS and SQL for device and user management systems.",
                "Developed scalable administrative services for device monitoring and management.",
                "Integrated AWS infrastructure to ensure reliable and scalable deployments.",
            ],
        },
        {
            "name": "ESEO – Sports Booking Platform",
            "bullets": [
                "Designed backend booking APIs using Node.js for sports facility booking workflows.",
                "Integrated Stripe payment gateway for secure online transactions.",
                "Provided support for React-based admin panel and production deployment workflows.",
            ],
        },
    ],
    "education": [
        {"degree": "B.Tech – Computer Science", "school": "Jaipur National University", "year": "2014–2018"},
        {"degree": "Higher Secondary (Science)", "school": "Paramount Academy", "year": "2012–2014"},
    ],
}

RESUME_TEXT = json.dumps(BASE_RESUME, indent=2)

# ── Claude API ────────────────────────────────────────────────────────────────

API_URL = "https://api.anthropic.com/v1/messages"
MODEL   = "claude-sonnet-4-20250514"

SYSTEM_PROMPT = """You are an expert resume writer. Given a base resume and a job description,
return a tailored version of the resume as a JSON object — same structure as the input.

Rules:
- Rewrite the summary (3-4 sentences) to reflect the JD's priorities and language
- Reorder skills to put JD-matched skills first; add any relevant skills Arjun has that match JD keywords
- Rewrite experience bullets to emphasise JD-relevant achievements; keep metrics intact
- Reorder project bullets by relevance to JD
- Keep all factual info accurate — do NOT invent new experience or skills
- Mirror the JD's tone (technical/startup/enterprise)
- Return ONLY valid JSON, no preamble, no markdown fences"""

def tailor_resume(jd_text: str, company: str) -> dict:
    """Call Claude API and return tailored resume as dict."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise EnvironmentError("ANTHROPIC_API_KEY not set. Run: export ANTHROPIC_API_KEY=sk-ant-...")

    prompt = f"""Base resume (JSON):
{RESUME_TEXT}

Job description for {company}:
{jd_text[:3000]}

Return the tailored resume as JSON only."""

    resp = requests.post(
        API_URL,
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": MODEL,
            "max_tokens": 2000,
            "system": SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=60,
    )
    resp.raise_for_status()
    raw = resp.json()["content"][0]["text"]
    raw = re.sub(r"^```json\s*|^```\s*|```$", "", raw.strip(), flags=re.MULTILINE).strip()
    return json.loads(raw)

# ── DOCX builder ──────────────────────────────────────────────────────────────

ACCENT   = RGBColor(0x18, 0x5F, 0xA5)   # blue
DARK     = RGBColor(0x1A, 0x1A, 0x1A)
MUTED    = RGBColor(0x55, 0x55, 0x55)
DIVIDER  = RGBColor(0xCC, 0xCC, 0xCC)

def _add_para_border_bottom(para, color_hex="CCCCCC", size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def build_docx(resume: dict, out_path: str):
    doc = Document()

    # Page margins — narrow for resume
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    def heading(text, size=11, color=ACCENT, bold=True, space_before=120, space_after=40, border=True):
        p = doc.add_paragraph()
        _set_spacing(p, before=space_before, after=space_after)
        if border:
            _add_para_border_bottom(p, "185FA5", 8)
        run = p.add_run(text.upper())
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Arial"
        return p

    def bullet(text, indent=0.25):
        p = doc.add_paragraph(style="List Bullet")
        _set_spacing(p, before=20, after=20)
        p.paragraph_format.left_indent  = Inches(indent)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = "Arial"
        run.font.color.rgb = DARK
        return p

    c = resume.get("contact", BASE_RESUME["contact"])

    # ── Name & title ──
    name_p = doc.add_paragraph()
    _set_spacing(name_p, before=0, after=20)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(resume.get("name", BASE_RESUME["name"]))
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT; r.font.name = "Arial"

    title_p = doc.add_paragraph()
    _set_spacing(title_p, before=0, after=30)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(resume.get("title", BASE_RESUME["title"]))
    r.font.size = Pt(11); r.font.color.rgb = MUTED; r.font.name = "Arial"

    contact_p = doc.add_paragraph()
    _set_spacing(contact_p, before=0, after=60)
    _add_para_border_bottom(contact_p, "185FA5", 12)
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [c.get("location",""), c.get("phone",""), c.get("email",""), c.get("linkedin","")]
    r = contact_p.add_run("  |  ".join(b for b in bits if b))
    r.font.size = Pt(9); r.font.color.rgb = MUTED; r.font.name = "Arial"

    # ── Summary ──
    heading("Professional Summary")
    s = resume.get("summary", BASE_RESUME["summary"])
    p = doc.add_paragraph()
    _set_spacing(p, before=40, after=40, line=276)
    r = p.add_run(s)
    r.font.size = Pt(9.5); r.font.name = "Arial"

    # ── Skills ──
    heading("Technical Skills")
    skills = resume.get("skills", BASE_RESUME["skills"])
    for cat, vals in skills.items():
        p = doc.add_paragraph()
        _set_spacing(p, before=20, after=20)
        r = p.add_run(f"{cat}: ")
        r.bold = True; r.font.size = Pt(9.5); r.font.name = "Arial"
        r2 = p.add_run(vals)
        r2.font.size = Pt(9.5); r2.font.name = "Arial"

    # ── Experience ──
    heading("Professional Experience")
    for job in resume.get("experience", BASE_RESUME["experience"]):
        p = doc.add_paragraph()
        _set_spacing(p, before=60, after=10)
        r = p.add_run(f"{job['role']}  |  {job['company']}")
        r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
        r2 = p.add_run(f"   {job['period']}")
        r2.font.size = Pt(9); r2.font.color.rgb = MUTED; r2.font.name = "Arial"
        for b in job.get("bullets", []):
            bullet(b)

    # ── Projects ──
    heading("Key Projects")
    for proj in resume.get("projects", BASE_RESUME["projects"]):
        p = doc.add_paragraph()
        _set_spacing(p, before=50, after=10)
        r = p.add_run(proj["name"])
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACCENT; r.font.name = "Arial"
        for b in proj.get("bullets", []):
            bullet(b)

    # ── Education ──
    heading("Education")
    for edu in resume.get("education", BASE_RESUME["education"]):
        p = doc.add_paragraph()
        _set_spacing(p, before=40, after=20)
        r = p.add_run(f"{edu['degree']}  |  {edu['school']}  |  {edu['year']}")
        r.font.size = Pt(9.5); r.font.name = "Arial"

    doc.save(out_path)

# ── Main pipeline ─────────────────────────────────────────────────────────────

def safe_filename(company: str, title: str) -> str:
    name = f"{company}_{title}"
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", name)[:80] + ".docx"

def main():
    parser = argparse.ArgumentParser(description="Generate JD-tailored resumes for Arjun Kumar")
    parser.add_argument("--csv",       default="arjun_jobs_top.csv", help="JobSpy CSV file")
    parser.add_argument("--limit",     type=int, default=0,          help="Max jobs to process (0 = all)")
    parser.add_argument("--min-score", type=int, default=50,         help="Skip jobs below this match score")
    parser.add_argument("--out-dir",   default="output_resumes",     help="Output directory")
    parser.add_argument("--delay",     type=float, default=2.0,      help="Seconds between API calls")
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(exist_ok=True)

    # Load CSV
    if not Path(args.csv).exists():
        print(f"\n  ERROR: '{args.csv}' not found.")
        print(f"  Run arjun_jobspy.py first to generate it.\n")
        return

    df = pd.read_csv(args.csv)
    df = df[df["match_score"] >= args.min_score].sort_values("match_score", ascending=False)
    if args.limit:
        df = df.head(args.limit)

    total = len(df)
    print(f"\n{'─'*58}")
    print(f"  Arjun Kumar — Resume Generator  |  {datetime.today().date()}")
    print(f"  Jobs to process : {total}  (score >= {args.min_score})")
    print(f"  Output folder   : {out_dir}/")
    print(f"{'─'*58}\n")

    summary_rows = []

    for i, (_, row) in enumerate(df.iterrows(), 1):
        company = str(row.get("company", "Unknown")).strip()
        title   = str(row.get("title",   "Role")).strip()
        jd_text = str(row.get("description", "")).strip()
        score   = int(row.get("match_score", 0))
        url     = str(row.get("job_url", ""))

        print(f"  [{i}/{total}] {company} — {title}  (score: {score}%)")

        filename = safe_filename(company, title)
        out_path = str(out_dir / filename)

        try:
            if jd_text and len(jd_text) > 100:
                tailored = tailor_resume(jd_text, company)
            else:
                print(f"    No JD text — using base resume")
                tailored = BASE_RESUME

            build_docx(tailored, out_path)
            print(f"    Saved → {filename}")
            summary_rows.append({"score": score, "company": company, "title": title, "file": filename, "url": url, "status": "ok"})

        except json.JSONDecodeError as e:
            print(f"    JSON parse error — saving base resume  ({e})")
            build_docx(BASE_RESUME, out_path)
            summary_rows.append({"score": score, "company": company, "title": title, "file": filename, "url": url, "status": "base_used"})

        except Exception as e:
            print(f"    ERROR: {e}")
            summary_rows.append({"score": score, "company": company, "title": title, "file": "—", "url": url, "status": f"error: {e}"})

        if i < total:
            time.sleep(args.delay)

    # Write summary CSV
    summary_path = out_dir / "summary.csv"
    pd.DataFrame(summary_rows).to_csv(str(summary_path), index=False)

    ok = sum(1 for r in summary_rows if r["status"] == "ok")
    print(f"\n{'─'*58}")
    print(f"  Done! {ok}/{total} resumes generated successfully.")
    print(f"  Summary → {summary_path}")
    print(f"{'─'*58}\n")

if __name__ == "__main__":
    main()